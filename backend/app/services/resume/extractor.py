"""
Document text extraction service for PDF and DOCX files.
"""
import os
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
import docx

from app.core.exceptions import FileProcessingError, FileValidationError


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF document using PyMuPDF (fitz)."""
    text_chunks = []
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text("text")
                if text:
                    text_chunks.append(text)
        return "\n".join(text_chunks).strip()
    except Exception as e:
        raise FileProcessingError(f"Failed to extract text from PDF file: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a Word DOCX document."""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as e:
        raise FileProcessingError(f"Failed to extract text from DOCX file: {str(e)}")


def extract_text_from_file(file_path: str, file_type: Optional[str] = None) -> str:
    """
    Extract raw text from a document based on its extension or specified file_type.
    """
    if not os.path.exists(file_path):
        raise FileValidationError(f"File not found on disk: {file_path}")

    ext = Path(file_path).suffix.lower()
    if file_type:
        file_type = file_type.lower()

    if ext == ".pdf" or (file_type and "pdf" in file_type):
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"] or (file_type and ("word" in file_type or "docx" in file_type)):
        return extract_text_from_docx(file_path)
    elif ext in [".txt"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        raise FileValidationError(f"Unsupported file format '{ext}'. Only PDF, DOCX, and TXT are supported.")
